"""
tool.py —— 工具调用执行层。

接收 Bot.message() 返回的 tool_calls 列表，依次执行每个工具，
将结果通过 Bot.add_tool_result() 写回历史，最终返回是否 FINISH。

替换/编辑模式的"两步流程"通过 working_config.json 的状态机维护，
与原版逻辑保持一致。
"""

import json
import traceback
from script.logger import log, user_log
from config import get_config
from script.system import system_command, find_file, edit_file
import os


# ── 单个工具执行 ──────────────────────────────────────────────────────────

def _execute_tool(name: str, args: dict,
                  input_func=input) -> tuple[str, dict[str, str], bool]:
    """执行单个工具调用。

    Returns:
        (result_str, file_contents_dict, is_finish)
        file_contents_dict: 仅 read_file 成功时非空，格式 {filename: content}
        is_finish:          True 表示 Bot 调用了 finish()
    """
    cfg = get_config()
    default_encoding: str = cfg.get('encoding', 'utf-8')

    match name:

        # ── set_wait ────────────────────────────────────────────────────
        case 'set_wait':
            from config import set_wait
            seconds = int(args.get('seconds', 10))
            set_wait(seconds)
            user_log(f'The timeout duration has been set to: {seconds} seconds', role="SETTINGS")
            return f'超时时长已更新为 {seconds} 秒', {}, False

        # ── set_read_limits ─────────────────────────────────────────────
        case 'set_read_limits':
            from config import set_read_limits
            MAX_LINES = 50000
            MAX_SIZE_KB = 5120
            raw_lines = args.get('max_lines')
            raw_size = args.get('max_size_kb')
            max_lines = min(int(raw_lines), MAX_LINES) if raw_lines is not None else None
            max_size_kb = min(int(raw_size), MAX_SIZE_KB) if raw_size is not None else None
            set_read_limits(max_lines=max_lines, max_size_kb=max_size_kb)
            parts = []
            if max_lines is not None:
                clamped = raw_lines != max_lines
                parts.append(f'Max lines: {max_lines}' + (' (clamped to limit)' if clamped else ''))
            if max_size_kb is not None:
                clamped = raw_size != max_size_kb
                parts.append(f'Max size: {max_size_kb} KB' + (' (clamped to limit)' if clamped else ''))
            msg = 'File read limits updated → ' + ', '.join(parts) if parts else 'No changes made'
            user_log(msg, role="SETTINGS")
            return msg, {}, False

        # ── finish ──────────────────────────────────────────────────────
        case 'finish':
            return 'FINISH', {}, True

        # ── system_command ──────────────────────────────────────────────
        case 'system_command':
            command = args.get('command', '')
            inputs = args.get('inputs')  # 获取新参数
            user_log(f'Shell Input: {command}{f' | input={inputs.split()}' if inputs is not None else ""}\n', role='SHELL')
            output = system_command(command, inputs=inputs)
            user_log(f'Shell Output: {"(NULL)" if output == "" else ("\n" + output)}', role='SHELL')
            return output or '（输出为空）', {}, False

        # ── edit_file ────────────────────────────────────────────────────
        case 'edit_file':
            file_path = args.get('file_path', '')
            content = args.get('content', '')
            encoding = args.get('encoding') or default_encoding
            try:
                edit_file(file_path, content, encoding)
                write_lines = len(content.splitlines())
                user_log(f'Write File: {file_path} (+{write_lines})')
                return f'文件写入完成: {file_path}（+{write_lines} 行）', {}, False
            except Exception as e:
                log(f'edit_file error | {file_path}\n{traceback.format_exc()}')
                return f'在编辑文件时遇到了以下错误: \n{type(e).__name__}: {e}\n可以尝试使用文件的绝对路径。', {}, False

        # ── replace_file ──────────────────────────────────────────────────
        case 'replace_file':
            file_path = args.get('file_path', '')
            old_text = args.get('old_text', '')
            new_text = args.get('new_text', '')
            encoding = args.get('encoding') or default_encoding
            try:
                content = find_file(file_path, encoding)
                if old_text not in content:
                    return f'替换失败: 在 {file_path} 中未找到指定的旧文本。', {}, False
                new_content = content.replace(old_text, new_text, 1)
                edit_file(file_path, new_content, encoding)
                old_lines = len(old_text.splitlines())
                new_lines = len(new_text.splitlines())
                user_log(f'Edit File: {file_path} (-{old_lines},+{new_lines})')
                return f'文件替换完成: {file_path}（-{old_lines} 行，+{new_lines} 行）', {}, False
            except Exception as e:
                log(f'replace_file error | {file_path}\n{traceback.format_exc()}')
                return f'在替换文件时遇到了以下错误: \n{type(e).__name__}: {e}\n可以尝试使用文件的绝对路径。', {}, False

        # ── read_file ──────────────────────────────────────────────────────
        case 'read_file':
            file_path = args.get('file_path', '')
            encoding = args.get('encoding') or default_encoding
            mode = args.get('mode', '')
            user_log(f'Read File: {file_path}' + (' [doc]' if mode == 'doc' else ''))
            try:
                import os as _os
                max_size_kb = cfg.get('read_max_size_kb', 100)
                max_lines = cfg.get('read_max_lines', 1000)
                max_size_bytes = max_size_kb * 1024
                file_size = _os.path.getsize(file_path)
                if file_size > max_size_bytes:
                    kb = file_size / 1024
                    return (f'文件过大: {file_path}（{kb:.1f} KB，当前体积限制: {max_size_kb} KB）。'
                            f'你可以使用 set_read_limits 调高限制，或使用其他方式读取部分内容。'), {}, False

                # ── doc 模式：用 python-docx 提取结构化 MD ──────────────
                if mode == 'doc':
                    import tempfile as _tempfile
                    import shutil as _shutil
                    from docx import Document as _Document
                    from docx.oxml.ns import qn as _qn

                    _ext = _os.path.splitext(file_path)[1].lower()
                    _tmp_dir = None
                    _docx_path = file_path

                    # 非 .docx 格式先用 LibreOffice 转换
                    if _ext != '.docx':
                        _lo = _shutil.which('libreoffice') or _shutil.which('soffice')
                        if not _lo:
                            return (
                                f'读取 {_ext} 格式需要 LibreOffice，但系统中未找到 libreoffice / soffice 命令。'
                            ), {}, False
                        _tmp_dir = _tempfile.mkdtemp()
                        try:
                            import subprocess as _sp
                            _cp = _sp.run(
                                [_lo, '--headless', '--convert-to', 'docx',
                                 '--outdir', _tmp_dir, file_path],
                                capture_output=True, timeout=30
                            )
                            if _cp.returncode != 0:
                                _err = _cp.stderr.decode(default_encoding, errors='replace').strip()
                                return (
                                    f'LibreOffice 转换失败（returncode={_cp.returncode}）: {_err}'
                                ), {}, False
                            _converted = [f for f in _os.listdir(_tmp_dir) if f.endswith('.docx')]
                            if not _converted:
                                return 'LibreOffice 转换完成但未生成 .docx 文件，转换可能不受支持。', {}, False
                            _docx_path = _os.path.join(_tmp_dir, _converted[0])
                        except _sp.TimeoutExpired:
                            return 'LibreOffice 转换超时（超过 30 秒）。', {}, False
                        except Exception as _e:
                            return f'调用 LibreOffice 时出错: {type(_e).__name__}: {_e}', {}, False

                    try:
                        doc = _Document(_docx_path)
                    finally:
                        if _tmp_dir:
                            _shutil.rmtree(_tmp_dir, ignore_errors=True)
                    lines = []
                    for para in doc.paragraphs:
                        style = para.style.name if para.style else ''
                        text = para.text
                        if style.startswith('Heading 1') or style.startswith('标题 1'):
                            lines.append(f'# {text}')
                        elif style.startswith('Heading 2') or style.startswith('标题 2'):
                            lines.append(f'## {text}')
                        elif style.startswith('Heading 3') or style.startswith('标题 3'):
                            lines.append(f'### {text}')
                        elif style.startswith('Heading') or style.startswith('标题'):
                            lines.append(f'#### {text}')
                        elif style.startswith('List Bullet') or style.startswith('列表段落'):
                            lines.append(f'- {text}')
                        elif style.startswith('List Number'):
                            lines.append(f'1. {text}')
                        else:
                            lines.append(text)
                    # 提取表格
                    for table in doc.tables:
                        rows = []
                        for i, row in enumerate(table.rows):
                            cells = [cell.text.replace('\n', ' ') for cell in row.cells]
                            rows.append('| ' + ' | '.join(cells) + ' |')
                            if i == 0:
                                rows.append('|' + '---|' * len(cells))
                        lines.append('')
                        lines.extend(rows)
                        lines.append('')
                    md_content = '\n'.join(lines)
                    line_count = len(md_content.splitlines())
                    if line_count > max_lines:
                        return (f'文件过大: {file_path}（提取后共 {line_count} 行，当前行数限制: {max_lines} 行）。'
                                f'可以使用 set_read_limits 调高限制。'), {}, False
                    result = f'打开文件（doc模式）: {file_path}\n{file_path}:\n{md_content}'
                    return result, {file_path: md_content}, False

                # ── 普通模式 ─────────────────────────────────────────────
                content = find_file(file_path, encoding)
                line_count = len(content.splitlines())
                if line_count > max_lines:
                    return (f'文件过大: {file_path}（共 {line_count} 行，当前行数限制: {max_lines} 行）。'
                            f'可以使用 set_read_limits 调高限制，或使用其他方式读取部分内容。'), {}, False
                result = f'打开文件: {file_path}\n{file_path}:\n{content}'
                return result, {file_path: content}, False
            except Exception as e:
                log(f'read_file error | {file_path}\n{traceback.format_exc()}')
                return f'在阅读文件时遇到了以下错误: \n{type(e).__name__}: {e}\n可以尝试使用文件的绝对路径。', {}, False

        # ── read_sheet ─────────────────────────────────────────────────────
        case 'read_sheet':
            import openpyxl as _openpyxl
            from openpyxl.utils import get_column_letter as _gcl
            file_path = args.get('file_path', '')
            sheet_name = args.get('sheet_name', '').strip()
            sheet_mode = args.get('sheet_mode', 'all')   # all / csv_only / formula_only
            range_str = args.get('range', '').strip()     # 如 A1:D20，空则自动检测
            user_log(f'Read Sheet: {file_path}' + (f' [{sheet_name}]' if sheet_name else ''))
            try:
                import os as _os
                import shutil as _shutil
                max_lines = cfg.get('read_max_lines', 1000)
                _sheet_ext = _os.path.splitext(file_path)[1].lower()

                # ── CSV 分支：用标准库直接读取 ──────────────────────────
                if _sheet_ext == '.csv':
                    import csv as _csv
                    with open(file_path, 'r', encoding=encoding, newline='') as _fh:
                        reader = _csv.reader(_fh)
                        csv_lines = []
                        for i, row in enumerate(reader):
                            if i >= max_lines:
                                csv_lines.append(f'（已截断，超过行数限制 {max_lines}，可用 set_read_limits 调高）')
                                break
                            csv_lines.append(','.join(row))
                    parts = [f'文件: {file_path}', f'格式: CSV', '']
                    parts.extend(csv_lines)
                    result = '\n'.join(parts)
                    return result, {file_path: result}, False

                # ── 非 .xlsx 系格式：用 LibreOffice 转换 ────────────────
                _xlsx_exts = {'.xlsx', '.xlsm', '.xltx', '.xltm'}
                _tmp_dir = None
                _xlsx_path = file_path
                if _sheet_ext not in _xlsx_exts:
                    _lo = _shutil.which('libreoffice') or _shutil.which('soffice')
                    if not _lo:
                        return (
                            f'读取 {_sheet_ext} 格式需要 LibreOffice，但系统中未找到 libreoffice / soffice 命令。'
                            f'请安装 LibreOffice 后重试，或手动将文件另存为 .xlsx 格式。'
                        ), {}, False
                    import tempfile as _tempfile, subprocess as _sp
                    _tmp_dir = _tempfile.mkdtemp()
                    try:
                        _cp = _sp.run(
                            [_lo, '--headless', '--convert-to', 'xlsx',
                             '--outdir', _tmp_dir, file_path],
                            capture_output=True, timeout=30
                        )
                        if _cp.returncode != 0:
                            _err = _cp.stderr.decode(default_encoding, errors='replace').strip()
                            _shutil.rmtree(_tmp_dir, ignore_errors=True)
                            return (
                                f'LibreOffice 转换失败（returncode={_cp.returncode}）: {_err}'
                            ), {}, False
                        _converted = [f for f in _os.listdir(_tmp_dir) if f.endswith('.xlsx')]
                        if not _converted:
                            _shutil.rmtree(_tmp_dir, ignore_errors=True)
                            return 'LibreOffice 转换完成但未生成 .xlsx 文件，转换可能不受支持。', {}, False
                        _xlsx_path = _os.path.join(_tmp_dir, _converted[0])
                    except _sp.TimeoutExpired:
                        _shutil.rmtree(_tmp_dir, ignore_errors=True)
                        return 'LibreOffice 转换超时（超过 30 秒）。', {}, False
                    except Exception as _e:
                        _shutil.rmtree(_tmp_dir, ignore_errors=True)
                        return f'调用 LibreOffice 时出错: {type(_e).__name__}: {_e}', {}, False

                try:
                    # 加载两次：一次取值，一次取公式
                    wb_val = _openpyxl.load_workbook(_xlsx_path, data_only=True)
                    wb_fml = _openpyxl.load_workbook(_xlsx_path, data_only=False)
                finally:
                    if _tmp_dir:
                        _shutil.rmtree(_tmp_dir, ignore_errors=True)

                sheet_names = wb_val.sheetnames

                # 未指定 sheet → 返回列表 + 第一个 sheet
                target_name = sheet_name if sheet_name and sheet_name in sheet_names else sheet_names[0]
                ws_val = wb_val[target_name]
                ws_fml = wb_fml[target_name]

                # 解析范围
                if range_str:
                    try:
                        min_col, min_row, max_col, max_row = _openpyxl.utils.cell.range_boundaries(range_str)
                    except Exception:
                        return f'无效的范围格式: {range_str}，请使用如 A1:D20 的格式。', {}, False
                else:
                    min_row, min_col = 1, 1
                    max_row = ws_val.max_row or 1
                    max_col = ws_val.max_column or 1

                # 限制行列数
                actual_rows = min(max_row - min_row + 1, max_lines)
                actual_cols = min(max_col - min_col + 1, max_lines)
                truncated = (actual_rows < max_row - min_row + 1) or (actual_cols < max_col - min_col + 1)
                max_row = min_row + actual_rows - 1
                max_col = min_col + actual_cols - 1

                # 构建 CSV 部分
                csv_lines = []
                formula_cells = []
                for r in range(min_row, max_row + 1):
                    row_vals = []
                    for c in range(min_col, max_col + 1):
                        cell_val = ws_val.cell(r, c)
                        cell_fml = ws_fml.cell(r, c)
                        fml = cell_fml.value
                        val = cell_val.value
                        row_vals.append('' if val is None else str(val))
                        # 收集公式单元格
                        if isinstance(fml, str) and fml.startswith('='):
                            coord = f'{_gcl(c)}{r}'
                            formula_cells.append(f'{coord}={fml}')
                    csv_lines.append(','.join(row_vals))

                # 组装结果
                parts = []
                parts.append(f'文件: {file_path}')
                parts.append(f'所有Sheet: {", ".join(sheet_names)}')
                parts.append(f'当前Sheet: {target_name}')
                if range_str:
                    parts.append(f'范围: {range_str}')
                if truncated:
                    parts.append(f'（已截断至 {actual_rows} 行 × {actual_cols} 列，当前限制: {max_lines}，可用 set_read_limits 调高）')
                parts.append('')

                if sheet_mode in ('all', 'csv_only'):
                    parts.append(f'[{target_name}] CSV:')
                    parts.extend(csv_lines)

                if sheet_mode in ('all', 'formula_only'):
                    parts.append('')
                    if formula_cells:
                        parts.append(f'[{target_name}] 公式单元格:')
                        parts.extend(formula_cells)
                    else:
                        parts.append(f'[{target_name}] 公式单元格: （无）')

                result = '\n'.join(parts)
                file_key = f'{file_path}::{target_name}'
                return result, {file_key: result}, False
            except Exception as e:
                log(f'read_sheet error | {file_path}\n{traceback.format_exc()}')
                return f'在阅读表格时遇到了以下错误: \n{type(e).__name__}: {e}\n可以尝试使用文件的绝对路径。', {}, False

        # ── get_skill ──────────────────────────────────────────────────────
        case 'get_skill':
            import os as _os
            skill_name = args.get('skill_name', '').strip()
            resource = args.get('resource', '').strip()  # 可选：scripts/xxx.py / references/xxx.md 等
            skills_dir = cfg.get('skills_dir', 'skill')
            if _os.path.isabs(skills_dir):
                skills_root = skills_dir
            else:
                _project_root = _os.environ.get('MOMOKA_PROJECT_DIR', '')
                skills_root = _os.path.join(_project_root, skills_dir)
            skill_path = _os.path.join(skills_root, skill_name)

            if not _os.path.isdir(skill_path):
                return (f'未找到skill: {skill_name}（路径: {skill_path}）'
                        f'\n提示: 请确认 {skills_dir}/ 目录下存在该skill文件夹。'), {}, False

            # 读取指定资源文件，或默认读取 SKILL.md
            if resource:
                target = _os.path.join(skill_path, resource)
                if not _os.path.isfile(target):
                    # 列出可用资源供模型参考
                    available = []
                    for root_, dirs_, files_ in _os.walk(skill_path):
                        for f_ in files_:
                            available.append(_os.path.join(root_, f_))
                    return (f'未找到资源文件: {resource}\n'
                            f'skill {skill_name!r} 中可用文件:\n' +
                            '\n'.join(f'  {f}' for f in sorted(available))), {}, False
                try:
                    with open(target, 'r', encoding=default_encoding) as fh:
                        content = fh.read()
                    user_log(f'Read Skill: {skill_name}/{resource}')
                    log(f'get_skill | {skill_name}/{resource} ({len(content)} chars)')
                    return content, {target: content}, False
                except Exception as e:
                    return f'读取资源文件失败: {e}', {}, False
            else:
                skill_md = _os.path.join(skill_path, 'SKILL.md')
                if not _os.path.isfile(skill_md):
                    return f'skill目录存在但缺少 SKILL.md: {skill_path}', {}, False
                try:
                    with open(skill_md, 'r', encoding=default_encoding) as fh:
                        content = fh.read()
                    # 顺带列出目录中可用的其他资源（scripts/references/assets）
                    extras = []
                    for sub in ('scripts', 'references', 'assets'):
                        sub_path = _os.path.join(skill_path, sub)
                        if _os.path.isdir(sub_path):
                            for fn in sorted(_os.listdir(sub_path)):
                                extras.append(os.path.realpath(os.path.join(sub_path, fn)))
                    suffix = ('\n\n可用资源文件（使用 resource 参数加载）:\n' +
                              '\n'.join(f'  {e}' for e in extras)) if extras else ''
                    user_log(f'Load Skill: {skill_name}')
                    log(f'get_skill | {skill_name}/SKILL.md ({len(content)} chars)')
                    return content + suffix, {skill_md: content}, False
                except Exception as e:
                    return f'读取 SKILL.md 失败: {e}', {}, False
        case 'change_directory':
            from script.system import set_cwd_explicit
            path = args.get('path', '')
            result = set_cwd_explicit(path)
            user_log(f'Switch Directory: {path}')
            return result, {}, False

        # ── ask_user ───────────────────────────────────────────────────────
        case 'ask_user':
            question = args.get('question', '')
            user_log(f'{question}', role='QUESTION')
            reply = input_func(f'>> ')
            return (f'用户回复: {reply}' if reply else '用户什么都没回复。'), {}, False

        # ── 浏览器指令 ─────────────────────────────────────────────────────

        case 'browse_open':
            from script.browser import browser_open
            url = args.get('url', '')
            user_log(f'Open Page: {url}', role="BROWSER")
            return browser_open(url), {}, False

        case 'browse_search':
            from script.browser import browser_search
            query = args.get('query', '')
            engine = args.get('engine', 'google')
            user_log(f'Search ({engine}): {query}', role='BROWSER')
            return browser_search(query, engine), {}, False

        case 'browse_read':
            from script.browser import browser_read, _page
            max_chars = args.get('max_chars', 4000)
            mode = args.get('mode', 'all')
            user_log(f'Reading Page ({mode})', role='BROWSER')
            result = browser_read(int(max_chars), mode)
            # 用 "browse_read:<url>" 作为 key，让 Bot 类折叠同页面的历史 read 结果
            try:
                current_url = _page.url if _page and not _page.is_closed() else None
            except Exception:
                current_url = None
            file_key = f'browse_read:{current_url}' if current_url else None
            file_contents = {file_key: result} if file_key else {}
            return result, file_contents, False

        case 'browse_click':
            from script.browser import browser_click
            element_uuid = args.get('element_uuid', '')
            user_log(f'Click: [{element_uuid}]', role='BROWSER')
            return browser_click(element_uuid), {}, False

        case 'browse_fill':
            from script.browser import browser_fill
            element_uuid = args.get('element_uuid', '')
            text = args.get('text', '')
            user_log(f'Fill: [{element_uuid}] → {text!r}', role='BROWSER')
            return browser_fill(element_uuid, text), {}, False

        case 'browse_press':
            from script.browser import browser_press
            element_uuid = args.get('element_uuid', '')
            key = args.get('key', 'Enter')
            user_log(f'Press: [{element_uuid}] {key!r}', role='BROWSER')
            return browser_press(element_uuid, key), {}, False

        case 'browse_find':
            from script.browser import browser_find
            text = args.get('text', '')
            max_results = args.get('max_results', 10)
            user_log(f'Page Search: {text!r}', role='BROWSER')
            return browser_find(text, int(max_results)), {}, False

        case 'browse_pdf':
            from script.browser import browser_pdf
            save_dir = args.get('save_dir') or cfg['work_dir']
            user_log(f'The webpage has been exported as a PDF and saved to: {save_dir}', role='BROWSER')
            return browser_pdf(save_dir), {}, False

        case 'browse_eval':
            from script.browser import browser_eval
            script = args.get('script', '')
            user_log(f'Eval: {script}', role='BROWSER')
            return browser_eval(script), {}, False

        case 'browse_wait_for_navigation':
            from script.browser import browser_wait_for_navigation
            timeout = args.get('timeout')
            state = args.get('state', 'networkidle')
            user_log(f'Loading ({state})...', role='BROWSER')
            return browser_wait_for_navigation(timeout, state), {}, False

        case 'browse_hover':
            from script.browser import browser_hover
            element_uuid = args.get('element_uuid', '')
            user_log(f'Hover: [{element_uuid}]', role='BROWSER')
            return browser_hover(element_uuid), {}, False

        case 'browse_select':
            from script.browser import browser_select
            element_uuid = args.get('element_uuid', '')
            value = args.get('value', '')
            user_log(f'Choose: [{element_uuid}] → {value!r}', role='BROWSER')
            return browser_select(element_uuid, value), {}, False

        case 'browse_get_url':
            from script.browser import browser_get_url
            user_log('Get current URL', role='BROWSER')
            return browser_get_url(), {}, False

        case 'browse_scroll':
            from script.browser import browser_scroll
            direction = args.get('direction', 'down')
            amount = int(args.get('amount', 500))
            element_uuid = args.get('element_uuid') or None
            user_log(f'Roll ({direction} {amount}px){"  [" + element_uuid + "]" if element_uuid else ""}', role='BROWSER')
            return browser_scroll(direction, amount, element_uuid), {}, False

        case 'browse_upload':
            from script.browser import browser_upload
            element_uuid = args.get('element_uuid', '')
            file_paths = args.get('file_paths', [])
            if isinstance(file_paths, str):
                file_paths = [file_paths]
            user_log(f'Upload File: [{element_uuid}] ← {file_paths}', role='BROWSER')
            return browser_upload(element_uuid, file_paths), {}, False

        case 'browse_download':
            from script.browser import browser_download
            element_uuid = args.get('element_uuid', '')
            save_dir = args.get('save_dir') or cfg.get('work_dir', '.')
            user_log(f'Download File: [{element_uuid}] → {save_dir}', role='BROWSER')
            return browser_download(element_uuid, save_dir), {}, False

        case 'browse_close':
            from script.browser import browser_close
            user_log('Close Browser', role='BROWSER')
            return browser_close(), {}, False

        case _:
            return f'未知工具: {name}', {}, False


# ── 主入口：处理一轮工具调用 ──────────────────────────────────────────────

def execute_tool_calls(
    work_bot,           # bot.Bot 实例
    tool_calls: list,
    input_func=input,
) -> tuple[bool, dict[str, str]]:
    """依次执行 tool_calls 列表中的所有工具，将结果写回 work_bot 历史。

    同时处理"文件编辑模式"和"替换模式"的多步状态机：
      - 进入编辑/替换模式后，后续 model 的纯文本输出即为文件内容或旧/新文本。

    Returns:
        (is_finish, all_file_contents)
        is_finish:          是否有工具调用了 finish()
        all_file_contents:  本轮所有 read_file 读到的文件内容合集
    """
    all_file_contents: dict[str, str] = {}
    is_finish = False

    for tc in tool_calls:
        name = tc.function.name
        try:
            args = json.loads(tc.function.arguments)
        except json.JSONDecodeError:
            args = {}

        result, file_contents, finish = _execute_tool(name, args, input_func)
        all_file_contents.update(file_contents)

        # browse_read：内容未变时用简短提示替换，避免全文重复进历史
        if name == 'browse_read' and file_contents:
            for file_key, new_content in file_contents.items():
                prev = next(
                    (m['file_contents'][file_key]
                     for m in reversed(work_bot._meta)
                     if file_key in m.get('file_contents', {})),
                    None
                )
                if prev is not None and prev == new_content:
                    url = file_key.removeprefix('browse_read:')
                    result = f'（页面内容与上次读取完全一致，无需重新解析。URL: {url}）'
                    log(f'execute_tool_calls | browse_read 内容未变化: {url}')

        log(f'execute_tool_calls | {name}({args}) → {result}')
        work_bot.add_tool_result(tc.id, result,
                                 file_contents=file_contents if file_contents else None)

        # 所有返回 file_contents 的工具，折叠历史中的旧版本，避免重复内容堆积
        if file_contents:
            for file_key in file_contents:
                collapsed = work_bot.collapse_file_in_history(file_key)
                if collapsed:
                    log(f'execute_tool_calls | 折叠历史 [{name}]: {file_key} ({collapsed} 条)')

        if finish:
            is_finish = True
            break

    return is_finish, all_file_contents