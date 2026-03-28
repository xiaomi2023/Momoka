"""
server/servers/office.py —— 文件与表格操作处理器。

涵盖: read_file / edit_file / replace_file / read_sheet
"""

from __future__ import annotations

import os
import traceback

from logger import log
from server import ToolResult, ToolContext


# ── edit_file ─────────────────────────────────────────────────────────────

def edit_file(args: dict, ctx: ToolContext) -> ToolResult:
    from server.servers.system import edit_file as _edit_file
    file_path = args.get('file_path', '')
    content = args.get('content', '')
    encoding = args.get('encoding') or ctx.cfg.get('encoding', 'utf-8')
    try:
        _edit_file(file_path, content, encoding)
        write_lines = len(content.splitlines())
        return ToolResult(
            text=f'文件写入完成: {file_path}（+{write_lines} 行）',
            log_msg=f'Write File: {file_path} (+{write_lines})',
        )
    except Exception as e:
        log(f'edit_file error | {file_path}\n{traceback.format_exc()}')
        return ToolResult(
            text=f'在编辑文件时遇到了以下错误: \n{type(e).__name__}: {e}\n可以尝试使用文件的绝对路径。'
        )


# ── replace_file ──────────────────────────────────────────────────────────

def replace_file(args: dict, ctx: ToolContext) -> ToolResult:
    from server.servers.system import find_file as _find_file, edit_file as _edit_file
    file_path = args.get('file_path', '')
    old_text = args.get('old_text', '')
    new_text = args.get('new_text', '')
    encoding = args.get('encoding') or ctx.cfg.get('encoding', 'utf-8')
    try:
        content = _find_file(file_path, encoding)
        if old_text not in content:
            return ToolResult(text=f'替换失败: 在 {file_path} 中未找到指定的旧文本。')
        new_content = content.replace(old_text, new_text, 1)
        _edit_file(file_path, new_content, encoding)
        old_lines = len(old_text.splitlines())
        new_lines = len(new_text.splitlines())
        return ToolResult(
            text=f'文件替换完成: {file_path}（-{old_lines} 行，+{new_lines} 行）',
            log_msg=f'Edit File: {file_path} (-{old_lines},+{new_lines})',
        )
    except Exception as e:
        log(f'replace_file error | {file_path}\n{traceback.format_exc()}')
        return ToolResult(
            text=f'在替换文件时遇到了以下错误: \n{type(e).__name__}: {e}\n可以尝试使用文件的绝对路径。'
        )


# ── read_file ─────────────────────────────────────────────────────────────

def read_file(args: dict, ctx: ToolContext) -> ToolResult:
    from server.servers.system import find_file as _find_file
    file_path = args.get('file_path', '')
    encoding = args.get('encoding') or ctx.cfg.get('encoding', 'utf-8')
    mode = args.get('mode', '')
    max_size_kb = ctx.cfg.get('read_max_size_kb', 100)
    max_lines = ctx.cfg.get('read_max_lines', 1000)

    log_label = f'Read File: {file_path}' + (' [doc]' if mode == 'doc' else '')

    try:
        max_size_bytes = max_size_kb * 1024
        file_size = os.path.getsize(file_path)
        if file_size > max_size_bytes:
            kb = file_size / 1024
            return ToolResult(
                text=(f'文件过大: {file_path}（{kb:.1f} KB，当前体积限制: {max_size_kb} KB）。'
                      f'你可以使用 set_read_limits 调高限制，或使用其他方式读取部分内容。'),
                log_msg=log_label,
            )

        # ── doc 模式 ─────────────────────────────────────────────────────
        if mode == 'doc':
            return _read_file_doc(file_path, encoding, max_lines, log_label, ctx)

        # ── 普通模式 ──────────────────────────────────────────────────────
        content = _find_file(file_path, encoding)
        line_count = len(content.splitlines())
        if line_count > max_lines:
            return ToolResult(
                text=(f'文件过大: {file_path}（共 {line_count} 行，当前行数限制: {max_lines} 行）。'
                      f'可以使用 set_read_limits 调高限制，或使用其他方式读取部分内容。'),
                log_msg=log_label,
            )
        return ToolResult(
            text=f'打开文件: {file_path}\n{file_path}:\n{content}',
            file_contents={file_path: content},
            log_msg=log_label,
        )

    except Exception as e:
        log(f'read_file error | {file_path}\n{traceback.format_exc()}')
        return ToolResult(
            text=f'在阅读文件时遇到了以下错误: \n{type(e).__name__}: {e}\n可以尝试使用文件的绝对路径。',
            log_msg=log_label,
        )


def _read_file_doc(file_path: str, encoding: str, max_lines: int,
                   log_label: str, ctx: ToolContext) -> ToolResult:
    """read_file 的 doc 模式：用 python-docx 提取结构化 Markdown。"""
    import shutil
    import tempfile
    import subprocess
    from docx import Document

    ext = os.path.splitext(file_path)[1].lower()
    tmp_dir = None
    docx_path = file_path

    if ext != '.docx':
        lo = shutil.which('libreoffice') or shutil.which('soffice')
        if not lo:
            return ToolResult(
                text=f'读取 {ext} 格式需要 LibreOffice，但系统中未找到 libreoffice / soffice 命令。',
                log_msg=log_label,
            )
        tmp_dir = tempfile.mkdtemp()
        try:
            cp = subprocess.run(
                [lo, '--headless', '--convert-to', 'docx', '--outdir', tmp_dir, file_path],
                capture_output=True, timeout=30,
            )
            if cp.returncode != 0:
                err = cp.stderr.decode(encoding, errors='replace').strip()
                return ToolResult(
                    text=f'LibreOffice 转换失败（returncode={cp.returncode}）: {err}',
                    log_msg=log_label,
                )
            converted = [f for f in os.listdir(tmp_dir) if f.endswith('.docx')]
            if not converted:
                return ToolResult(
                    text='LibreOffice 转换完成但未生成 .docx 文件，转换可能不受支持。',
                    log_msg=log_label,
                )
            docx_path = os.path.join(tmp_dir, converted[0])
        except subprocess.TimeoutExpired:
            return ToolResult(text='LibreOffice 转换超时（超过 30 秒）。', log_msg=log_label)
        except Exception as e:
            return ToolResult(text=f'调用 LibreOffice 时出错: {type(e).__name__}: {e}', log_msg=log_label)

    try:
        doc = Document(docx_path)
    finally:
        if tmp_dir:
            shutil.rmtree(tmp_dir, ignore_errors=True)

    lines = []
    _HEADING_MAP = {
        'Heading 1': '#', '标题 1': '#',
        'Heading 2': '##', '标题 2': '##',
        'Heading 3': '###', '标题 3': '###',
    }
    for para in doc.paragraphs:
        style = para.style.name if para.style else ''
        text = para.text
        prefix = next((v for k, v in _HEADING_MAP.items() if style.startswith(k)), None)
        if prefix:
            lines.append(f'{prefix} {text}')
        elif style.startswith('Heading') or style.startswith('标题'):
            lines.append(f'#### {text}')
        elif style.startswith('List Bullet') or style.startswith('列表段落'):
            lines.append(f'- {text}')
        elif style.startswith('List Number'):
            lines.append(f'1. {text}')
        else:
            lines.append(text)

    for table in doc.tables:
        rows = []
        for i, row in enumerate(table.rows):
            cells = [cell.text.replace('\n', ' ') for cell in row.cells]
            rows.append('| ' + ' | '.join(cells) + ' |')
            if i == 0:
                rows.append('|' + '---|' * len(cells))
        lines.append('')
        lines.extend(rows)
        lines.append('')

    md_content = '\n'.join(lines)
    line_count = len(md_content.splitlines())
    if line_count > max_lines:
        return ToolResult(
            text=(f'文件过大: {file_path}（提取后共 {line_count} 行，当前行数限制: {max_lines} 行）。'
                  f'可以使用 set_read_limits 调高限制。'),
            log_msg=log_label,
        )
    return ToolResult(
        text=f'打开文件（doc模式）: {file_path}\n{file_path}:\n{md_content}',
        file_contents={file_path: md_content},
        log_msg=log_label,
    )


# ── read_sheet ────────────────────────────────────────────────────────────

def read_sheet(args: dict, ctx: ToolContext) -> ToolResult:
    import openpyxl
    from openpyxl.utils import get_column_letter as gcl

    file_path = args.get('file_path', '')
    sheet_name = args.get('sheet_name', '').strip()
    sheet_mode = args.get('sheet_mode', 'all')
    range_str = args.get('range', '').strip()
    encoding = args.get('encoding') or ctx.cfg.get('encoding', 'utf-8')
    max_lines = ctx.cfg.get('read_max_lines', 1000)

    log_label = f'Read Sheet: {file_path}' + (f' [{sheet_name}]' if sheet_name else '')

    try:
        ext = os.path.splitext(file_path)[1].lower()

        # ── CSV ──────────────────────────────────────────────────────────
        if ext == '.csv':
            return _read_csv(file_path, encoding, max_lines, log_label)

        # ── 非 xlsx：先用 LibreOffice 转换 ───────────────────────────────
        xlsx_exts = {'.xlsx', '.xlsm', '.xltx', '.xltm'}
        xlsx_path, tmp_dir = _ensure_xlsx(file_path, ext, xlsx_exts, encoding, log_label)
        if isinstance(xlsx_path, ToolResult):   # 转换失败时直接返回错误
            return xlsx_path

        try:
            wb_val = openpyxl.load_workbook(xlsx_path, data_only=True)
            wb_fml = openpyxl.load_workbook(xlsx_path, data_only=False)
        finally:
            if tmp_dir:
                import shutil
                shutil.rmtree(tmp_dir, ignore_errors=True)

        sheet_names = wb_val.sheetnames
        target_name = sheet_name if sheet_name and sheet_name in sheet_names else sheet_names[0]
        ws_val = wb_val[target_name]
        ws_fml = wb_fml[target_name]

        # ── 范围解析 ──────────────────────────────────────────────────────
        if range_str:
            try:
                min_col, min_row, max_col, max_row = openpyxl.utils.cell.range_boundaries(range_str)
            except Exception:
                return ToolResult(text=f'无效的范围格式: {range_str}，请使用如 A1:D20 的格式。',
                                  log_msg=log_label)
        else:
            min_row, min_col = 1, 1
            max_row = ws_val.max_row or 1
            max_col = ws_val.max_column or 1

        actual_rows = min(max_row - min_row + 1, max_lines)
        actual_cols = min(max_col - min_col + 1, max_lines)
        truncated = (actual_rows < max_row - min_row + 1) or (actual_cols < max_col - min_col + 1)
        max_row = min_row + actual_rows - 1
        max_col = min_col + actual_cols - 1

        csv_lines = []
        formula_cells = []
        for r in range(min_row, max_row + 1):
            row_vals = []
            for c in range(min_col, max_col + 1):
                val = ws_val.cell(r, c).value
                fml = ws_fml.cell(r, c).value
                row_vals.append('' if val is None else str(val))
                if isinstance(fml, str) and fml.startswith('='):
                    formula_cells.append(f'{gcl(c)}{r}={fml}')
            csv_lines.append(','.join(row_vals))

        parts = [
            f'文件: {file_path}',
            f'所有Sheet: {", ".join(sheet_names)}',
            f'当前Sheet: {target_name}',
        ]
        if range_str:
            parts.append(f'范围: {range_str}')
        if truncated:
            parts.append(f'（已截断至 {actual_rows} 行 × {actual_cols} 列，'
                         f'当前限制: {max_lines}，可用 set_read_limits 调高）')
        parts.append('')

        if sheet_mode in ('all', 'csv_only'):
            parts.append(f'[{target_name}] CSV:')
            parts.extend(csv_lines)

        if sheet_mode in ('all', 'formula_only'):
            parts.append('')
            if formula_cells:
                parts.append(f'[{target_name}] 公式单元格:')
                parts.extend(formula_cells)
            else:
                parts.append(f'[{target_name}] 公式单元格: （无）')

        result = '\n'.join(parts)
        file_key = f'{file_path}::{target_name}'
        return ToolResult(text=result, file_contents={file_key: result}, log_msg=log_label)

    except Exception as e:
        log(f'read_sheet error | {file_path}\n{traceback.format_exc()}')
        return ToolResult(
            text=f'在阅读表格时遇到了以下错误: \n{type(e).__name__}: {e}\n可以尝试使用文件的绝对路径。',
            log_msg=log_label,
        )


def _read_csv(file_path: str, encoding: str, max_lines: int, log_label: str) -> ToolResult:
    import csv
    with open(file_path, 'r', encoding=encoding, newline='') as fh:
        reader = csv.reader(fh)
        csv_lines = []
        for i, row in enumerate(reader):
            if i >= max_lines:
                csv_lines.append(f'（已截断，超过行数限制 {max_lines}，可用 set_read_limits 调高）')
                break
            csv_lines.append(','.join(row))
    parts = [f'文件: {file_path}', '格式: CSV', '']
    parts.extend(csv_lines)
    result = '\n'.join(parts)
    return ToolResult(text=result, file_contents={file_path: result}, log_msg=log_label)


def _ensure_xlsx(file_path: str, ext: str, xlsx_exts: set,
                 encoding: str, log_label: str):
    """若文件不是 xlsx 系列，调用 LibreOffice 转换；返回 (xlsx_path, tmp_dir)。
    转换失败时返回 (ToolResult, None)。
    """
    if ext in xlsx_exts:
        return file_path, None

    import shutil
    import tempfile
    import subprocess

    lo = shutil.which('libreoffice') or shutil.which('soffice')
    if not lo:
        return ToolResult(
            text=(f'读取 {ext} 格式需要 LibreOffice，但系统中未找到 libreoffice / soffice 命令。'
                  f'请安装 LibreOffice 后重试，或手动将文件另存为 .xlsx 格式。'),
            log_msg=log_label,
        ), None

    tmp_dir = tempfile.mkdtemp()
    try:
        cp = subprocess.run(
            [lo, '--headless', '--convert-to', 'xlsx', '--outdir', tmp_dir, file_path],
            capture_output=True, timeout=30,
        )
        if cp.returncode != 0:
            err = cp.stderr.decode(encoding, errors='replace').strip()
            shutil.rmtree(tmp_dir, ignore_errors=True)
            return ToolResult(
                text=f'LibreOffice 转换失败（returncode={cp.returncode}）: {err}',
                log_msg=log_label,
            ), None
        converted = [f for f in os.listdir(tmp_dir) if f.endswith('.xlsx')]
        if not converted:
            shutil.rmtree(tmp_dir, ignore_errors=True)
            return ToolResult(
                text='LibreOffice 转换完成但未生成 .xlsx 文件，转换可能不受支持。',
                log_msg=log_label,
            ), None
        return os.path.join(tmp_dir, converted[0]), tmp_dir
    except subprocess.TimeoutExpired:
        shutil.rmtree(tmp_dir, ignore_errors=True)
        return ToolResult(text='LibreOffice 转换超时（超过 30 秒）。', log_msg=log_label), None
    except Exception as e:
        shutil.rmtree(tmp_dir, ignore_errors=True)
        return ToolResult(
            text=f'调用 LibreOffice 时出错: {type(e).__name__}: {e}',
            log_msg=log_label,
        ), None