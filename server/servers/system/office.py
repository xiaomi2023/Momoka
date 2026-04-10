"""
server/servers/system/office.py —— Office file operation tools.

Covers: docx reading (doc mode), xlsx/csv reading, LibreOffice conversion.
"""

from __future__ import annotations

import csv
import os
import shutil
import subprocess
import tempfile

from logger import log
from server.types import ToolResult
from config import get_working_config

# ── 可选依赖检测 ─────────────────────────────────────────────────────────

try:
    from docx import Document
    _HAS_DOCX = True
except ImportError:
    _HAS_DOCX = False
    Document = None  # type: ignore

try:
    import openpyxl
    from openpyxl.utils import get_column_letter as gcl
    _HAS_OPENPYXL = True
except ImportError:
    _HAS_OPENPYXL = False
    openpyxl = None  # type: ignore
    gcl = None  # type: ignore


# ── Missing dependency messages ───────────────────────────────────────────

_MISSING_DOCX_MSG = (
    "Cannot read Word document: python-docx library is missing.\n"
    "Consider using other methods to read"
)

_MISSING_OPENPYXL_MSG = (
    "Cannot read Excel file: openpyxl library is missing.\n"
    "Consider reading CSV format files directly or using other methods"
)


# ── DOCX reading ───────────────────────────────────────────────────────────

_HEADING_MAP = {
    'Heading 1': '#', '标题 1': '#',
    'Heading 2': '##', '标题 2': '##',
    'Heading 3': '###', '标题 3': '###',
}


def read_docx(file_path: str, encoding: str, max_lines: int, log_label: str) -> ToolResult:
    """Extract Word document to Markdown format using python-docx."""
    if not _HAS_DOCX:
        return ToolResult(
            text=_MISSING_DOCX_MSG,
            log_msg=[
                (log_label, 'TOOL'),
                ('Missing dependency: python-docx is required to read Word documents. \nInstall with: pip install python-docx', 'WARN'),
            ],
        )

    tmp_dir = None
    docx_path = file_path
    ext = os.path.splitext(file_path)[1].lower()

    if ext != '.docx':
        lo = shutil.which('libreoffice') or shutil.which('soffice')
        if not lo:
            return ToolResult(
                text=f'<Reading {ext} format requires LibreOffice, but libreoffice / soffice command not found>',
                log_msg=log_label,
            )
        tmp_dir = tempfile.mkdtemp()
        timeout = get_working_config().get('wait', 10)
        try:
            cp = subprocess.run(
                [lo, '--headless', '--convert-to', 'docx', '--outdir', tmp_dir, file_path],
                capture_output=True, timeout=timeout,
            )
            if cp.returncode != 0:
                err = cp.stderr.decode(encoding, errors='replace').strip()
                return ToolResult(
                    text=f'<LibreOffice conversion failed (returncode={cp.returncode}): {err}>',
                    log_msg=log_label,
                )
            converted = [f for f in os.listdir(tmp_dir) if f.endswith('.docx')]
            if not converted:
                return ToolResult(
                    text='<LibreOffice conversion completed but no .docx file was generated, conversion may not be supported>',
                    log_msg=log_label,
                )
            docx_path = os.path.join(tmp_dir, converted[0])
        except subprocess.TimeoutExpired:
            timeout = get_working_config().get('wait', 10)
            return ToolResult(text=f'<LibreOffice conversion timed out (exceeded {timeout} seconds)>', log_msg=log_label)
        except Exception as e:
            return ToolResult(text=f'<Error calling LibreOffice: {type(e).__name__}: {e}>', log_msg=log_label)

    try:
        doc = Document(docx_path)  # type: ignore
    finally:
        if tmp_dir:
            shutil.rmtree(tmp_dir, ignore_errors=True)

    lines = []
    for para in doc.paragraphs:
        style = para.style.name if para.style else ''
        text = para.text
        prefix = next((v for k, v in _HEADING_MAP.items() if style.startswith(k)), None)
        if prefix:
            lines.append(f'{prefix} {text}')
        elif style.startswith('Heading') or style.startswith('标题'):
            lines.append(f'#### {text}')
        elif style.startswith('List Bullet') or style.startswith('列表段落'):
            lines.append(f'- {text}')
        elif style.startswith('List Number'):
            lines.append(f'1. {text}')
        else:
            lines.append(text)

    for table in doc.tables:
        rows = []
        for i, row in enumerate(table.rows):
            cells = [cell.text.replace('\n', ' ') for cell in row.cells]
            rows.append('| ' + ' | '.join(cells) + ' |')
            if i == 0:
                rows.append('|' + '---|' * len(cells))
        lines.append('')
        lines.extend(rows)
        lines.append('')

    md_content = '\n'.join(lines)
    line_count = len(md_content.splitlines())
    if line_count > max_lines:
        return ToolResult(
            text=(f'<File too large: {file_path} ({line_count} lines after extraction, current line limit: {max_lines}), '
                  f'Consider using set_read_limits to increase the limits>'),
            log_msg=log_label,
        )
    file_name = os.path.basename(file_path)
    return ToolResult(
        text=f'<Opened File {file_name} in {file_path}>\n'
             f'<{file_name}>\n'
             f'{md_content}\n'
             f'</{file_name}>',
        file_contents={file_path: md_content},
        log_msg=log_label,
    )


# ── Excel/CSV reading ─────────────────────────────────────────────────────

def read_sheet_tool(file_path: str, sheet_name: str, sheet_mode: str,
                    range_str: str, encoding: str, max_lines: int, log_label: str) -> ToolResult:
    """Read Excel or CSV file."""
    ext = os.path.splitext(file_path)[1].lower()

    # ── CSV ──────────────────────────────────────────────────────────
    if ext == '.csv':
        return _read_csv_file(file_path, encoding, max_lines, log_label)

    # ── Excel requires openpyxl ─────────────────────────────────────
    if not _HAS_OPENPYXL:
        return ToolResult(
            text=_MISSING_OPENPYXL_MSG,
            log_msg=[
                (log_label, 'TOOL'),
                ('Missing dependency: openpyxl is required to read Excel files. \nInstall with: pip install openpyxl', 'WARN'),
            ],
        )

    # ── Non-xlsx: convert with LibreOffice first ─────────────────────
    xlsx_exts = {'.xlsx', '.xlsm', '.xltx', '.xltm'}
    xlsx_path, tmp_dir = _ensure_xlsx(file_path, ext, xlsx_exts, encoding, log_label)
    if isinstance(xlsx_path, ToolResult):  # 转换失败时直接返回错误
        return xlsx_path

    try:
        wb_val = openpyxl.load_workbook(xlsx_path, data_only=True)  # type: ignore
        wb_fml = openpyxl.load_workbook(xlsx_path, data_only=False)  # type: ignore
    finally:
        if tmp_dir:
            shutil.rmtree(tmp_dir, ignore_errors=True)

    sheet_names = wb_val.sheetnames
    target_name = sheet_name if sheet_name and sheet_name in sheet_names else sheet_names[0]
    ws_val = wb_val[target_name]
    ws_fml = wb_fml[target_name]

    # ── Range parsing ────────────────────────────────────────────────
    if range_str:
        try:
            min_col, min_row, max_col, max_row = openpyxl.utils.cell.range_boundaries(range_str)  # type: ignore
        except Exception:
            return ToolResult(text=f'<Invalid range format: {range_str}, please use format like A1:D20>',
                              log_msg=log_label)
    else:
        min_row, min_col = 1, 1
        max_row = ws_val.max_row or 1
        max_col = ws_val.max_column or 1

    actual_rows = min(max_row - min_row + 1, max_lines)
    actual_cols = min(max_col - min_col + 1, max_lines)
    truncated = (actual_rows < max_row - min_row + 1) or (actual_cols < max_col - min_col + 1)
    max_row = min_row + actual_rows - 1
    max_col = min_col + actual_cols - 1

    csv_lines = []
    formula_cells = []
    for r in range(min_row, max_row + 1):
        row_vals = []
        for c in range(min_col, max_col + 1):
            val = ws_val.cell(r, c).value
            fml = ws_fml.cell(r, c).value
            row_vals.append('' if val is None else str(val))
            if isinstance(fml, str) and fml.startswith('='):
                formula_cells.append(f'{gcl(c)}{r}={fml}')  # type: ignore
        csv_lines.append(','.join(row_vals))

    parts = [
        f'File: {file_path}',
        f'All Sheets: {", ".join(sheet_names)}',
        f'Current Sheet: {target_name}',
    ]
    if range_str:
        parts.append(f'Range: {range_str}')
    if truncated:
        parts.append(f'(Truncated to {actual_rows} rows × {actual_cols} columns, '
                     f'current limit: {max_lines}, use set_read_limits to increase)')
    parts.append('')

    if sheet_mode in ('all', 'csv_only'):
        parts.append(f'[{target_name}] CSV:')
        parts.extend(csv_lines)

    if sheet_mode in ('all', 'formula_only'):
        parts.append('')
        if formula_cells:
            parts.append(f'[{target_name}] Formula Cells:')
            parts.extend(formula_cells)
        else:
            parts.append(f'[{target_name}] Formula Cells: (none)')

    result = '\n'.join(parts)
    file_key = f'{file_path}::{target_name}'
    return ToolResult(text=result, file_contents={file_key: result}, log_msg=log_label)


def _read_csv_file(file_path: str, encoding: str, max_lines: int, log_label: str) -> ToolResult:
    """Read CSV file."""
    with open(file_path, 'r', encoding=encoding, newline='') as fh:
        reader = csv.reader(fh)
        csv_lines = []
        for i, row in enumerate(reader):
            if i >= max_lines:
                csv_lines.append(f'(Truncated, exceeded line limit {max_lines}, use set_read_limits to increase)')
                break
            csv_lines.append(','.join(row))
    parts = [f'File: {file_path}', 'Format: CSV', '']
    parts.extend(csv_lines)
    result = '\n'.join(parts)
    return ToolResult(text=result, file_contents={file_path: result}, log_msg=log_label)


def _ensure_xlsx(file_path: str, ext: str, xlsx_exts: set,
                 encoding: str, log_label: str) -> tuple[str | ToolResult, str | None]:
    """If file is not xlsx format, convert with LibreOffice; returns (xlsx_path, tmp_dir).
    Returns (ToolResult, None) on conversion failure.
    """
    if ext in xlsx_exts:
        return file_path, None

    lo = shutil.which('libreoffice') or shutil.which('soffice')
    if not lo:
        return ToolResult(
            text=(f'<Reading {ext} format requires LibreOffice, but libreoffice / soffice command not found. '
                  f'Consider manually saving the file as .xlsx format>'),
            log_msg=log_label,
        ), None

    tmp_dir = tempfile.mkdtemp()
    try:
        timeout = get_working_config().get('wait', 10)
        cp = subprocess.run(
            [lo, '--headless', '--convert-to', 'xlsx', '--outdir', tmp_dir, file_path],
            capture_output=True, timeout=timeout,
        )
        if cp.returncode != 0:
            err = cp.stderr.decode(encoding, errors='replace').strip()
            shutil.rmtree(tmp_dir, ignore_errors=True)
            return ToolResult(
                text=f'<LibreOffice conversion failed (returncode={cp.returncode}): {err}>',
                log_msg=log_label,
            ), None
        converted = [f for f in os.listdir(tmp_dir) if f.endswith('.xlsx')]
        if not converted:
            shutil.rmtree(tmp_dir, ignore_errors=True)
            return ToolResult(
                text='<LibreOffice conversion completed but no .xlsx file was generated, conversion may not be supported>',
                log_msg=log_label,
            ), None
        return os.path.join(tmp_dir, converted[0]), tmp_dir
    except subprocess.TimeoutExpired:
        shutil.rmtree(tmp_dir, ignore_errors=True)
        timeout = get_working_config().get('wait', 10)
        return ToolResult(text=f'<LibreOffice conversion timed out (exceeded {timeout} seconds)>', log_msg=log_label), None
    except Exception as e:
        shutil.rmtree(tmp_dir, ignore_errors=True)
        return ToolResult(
            text=f'<Error calling LibreOffice: {type(e).__name__}: {e}>',
            log_msg=log_label,
        ), None
